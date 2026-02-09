"""
Intelligent Document Generation Agent
======================================
Takes a Process Analysis Agent session file (JSON) and uses OpenAI's GPT
to *comprehend* the conversation, extract structured data, and generate
a professional Word document with:
  - Executive Summary
  - SIPOC Diagram (table)
  - Detailed AS-IS Process Map (table)
  - Pain Points & Improvement Opportunities
  - Baseline Metrics
  - Mermaid.js Flowchart (embedded as text + optional image)

Usage:
    python intelligent_doc_generator.py <session_file.json>
    python intelligent_doc_generator.py outputs/sd_light_session.json

Requirements:
    pip install openai python-docx python-dotenv
"""

import json
import os
import subprocess
import sys
import textwrap
from datetime import datetime
from pathlib import Path
import shutil

from dotenv import load_dotenv
from openai import OpenAI

# ── Load environment ─────────────────────────────────────────────────
# Try multiple locations for .env
for env_candidate in [
    Path(__file__).parent / ".env",
    Path(__file__).parent.parent / ".env",
    Path.cwd() / ".env",
]:
    if env_candidate.exists():
        load_dotenv(dotenv_path=env_candidate)
        break

# ── Constants ────────────────────────────────────────────────────────
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")
MAX_TOKENS = 4096


# =====================================================================
#  STEP 1 — AI-powered extraction prompts
# =====================================================================

EXTRACTION_SYSTEM_PROMPT = textwrap.dedent("""\
You are a senior business process analyst. You will receive the full
conversation transcript of a process-analysis interview.

Your job is to read the ENTIRE conversation carefully and extract
structured data. Return ONLY valid JSON — no markdown fences, no
explanation, no preamble. Just the raw JSON object.
""")

SIPOC_PROMPT = textwrap.dedent("""\
From the conversation below, extract a complete SIPOC diagram.

Return JSON in EXACTLY this structure:
{
  "process_name": "...",
  "process_owner": "...",
  "frequency": "...",
  "org_position": "...",
  "suppliers": [
    {"name": "...", "description": "What they supply / their role"}
  ],
  "inputs": [
    {"name": "...", "description": "Details about this input, format, trigger, etc."}
  ],
  "process_steps_high_level": [
    "Step 1 description",
    "Step 2 description"
  ],
  "outputs": [
    {"name": "...", "description": "Details about this output"}
  ],
  "customers": [
    {"name": "...", "description": "Why they receive the output"}
  ],
  "systems_and_tools": ["System 1", "System 2"],
  "known_exceptions": ["Exception 1", "Exception 2"]
}

Be thorough. Include every detail mentioned in the conversation.
If something wasn't discussed, use "Not discussed" as the value.

CONVERSATION:
""")

PROCESS_MAP_PROMPT = textwrap.dedent("""\
From the conversation below, extract a DETAILED AS-IS process map.

For EVERY step mentioned (including sub-steps, decision points, and
exception paths), create an entry.

Return JSON in EXACTLY this structure:
{
  "steps": [
    {
      "step_number": 1,
      "step_name": "Short name",
      "description": "Detailed description of what happens",
      "performer": "Who does this (role/team)",
      "system_tool": "What system or tool is used",
      "input": "What triggers or feeds this step",
      "output": "What this step produces",
      "decision_point": true/false,
      "decision_options": ["Option A", "Option B"],
      "pain_point": "Any pain point or inefficiency mentioned for this step, or null",
      "estimated_time": "Time estimate if mentioned, or null",
      "notes": "Any additional context"
    }
  ],
  "handoffs": [
    {
      "from_role": "...",
      "to_role": "...",
      "what": "What is handed off",
      "how": "How (email, system, etc.)"
    }
  ],
  "pain_points_summary": [
    {
      "category": "Manual work / Waiting / Errors / etc.",
      "description": "Detailed description of the pain point",
      "impact": "What impact this has",
      "affected_steps": [1, 2]
    }
  ]
}

Be as detailed as possible. If the conversation mentions sub-steps,
include them. Every piece of process information matters.

CONVERSATION:
""")

BASELINE_PROMPT = textwrap.dedent("""\
From the conversation below, extract ALL baseline metrics and
quantitative information about the current process.

Return JSON in EXACTLY this structure:
{
  "volume": {
    "description": "How often the process runs and volume details",
    "frequency": "Daily / Weekly / Monthly / etc.",
    "volume_per_period": "Number of items processed per period",
    "details": "Any additional volume context"
  },
  "time": {
    "total_cycle_time": "End-to-end time",
    "active_work_time": "Actual hands-on time",
    "waiting_time": "Time spent waiting",
    "step_breakdown": [
      {"step": "Step name", "time": "Duration", "notes": "Context"}
    ]
  },
  "people": {
    "fte_involved": "Number of people / FTE",
    "roles": ["Role 1", "Role 2"],
    "time_investment_per_period": "Hours per week/month spent",
    "cost_estimate": "If mentioned"
  },
  "quality": {
    "error_rate": "Percentage or description of errors",
    "rework_rate": "How often things need to be redone",
    "common_errors": ["Error type 1", "Error type 2"]
  },
  "risk_and_compliance": {
    "risks": ["Risk 1", "Risk 2"],
    "sla": "Any SLA or deadline requirements",
    "compliance_requirements": "Any regulatory or compliance needs"
  },
  "executive_summary": "A 3-5 sentence executive summary of the process, its current state, key challenges, and potential for improvement. Write this as a professional analyst would for a steering committee."
}

Extract EVERY number, time estimate, frequency, and quantitative
detail from the conversation. If a metric wasn't discussed, set it
to "Not measured / Not discussed".

CONVERSATION:
""")

MERMAID_PROMPT = textwrap.dedent("""\
From the conversation below, create a Mermaid.js flowchart diagram
that accurately represents the AS-IS process.

LAYOUT AND STYLE REQUIREMENTS:
- Use `graph LR` (left-to-right / horizontal) layout
- Use subgraph blocks as HORIZONTAL SWIMLANES — one per role/team
- Decision points MUST use curly braces for diamond shape: `D1{Is it valid?}`
- Regular steps use square brackets: `A1[Step name]`
- Start/end nodes use rounded brackets: `S([Start])` and `E([End])`
- Mark pain points with `:::painPoint` class
- Include ALL exception paths and branches
- Use descriptive edge labels with `-->|label text|`
- Keep node labels short but clear (max 5-6 words per node)

CRITICAL FORMATTING RULES:
- Each node, arrow, subgraph, and end statement MUST be on its OWN LINE
- Use proper indentation (2 spaces inside subgraphs)
- NEVER put multiple statements on the same line
- Do NOT use special characters inside node labels or edge labels:
  Write EUR not €, GBP not £, USD not $
  Write "less than" not <, "greater than" not >
  Write "and" not &
- Always end with style definitions

EXAMPLE of correct format and style:
graph LR
  subgraph Requester
    S([Start]) --> A1[Submit request]
    A1 --> A2[Upload documents]
  end
  subgraph Processing Team
    A2 --> B1{Documents complete?}
    B1 -->|Yes| B2[Process request]
    B1 -->|No| B3[Return to requester]:::painPoint
    B3 --> A1
    B2 --> B4{Approved?}
    B4 -->|Yes| C1[Send confirmation]
    B4 -->|No| B5[Escalate]:::painPoint
  end
  subgraph Customer
    C1 --> D1[Receive output]
    D1 --> E([End])
  end
  classDef painPoint fill:#ff6b6b,stroke:#c92a2a,stroke-width:2px,color:#fff
  style S fill:#51cf66,stroke:#2b8a3e,color:#fff
  style E fill:#868e96,stroke:#495057,color:#fff

Return ONLY the Mermaid code. No JSON, no explanation, no markdown
fences, no wrapping. Start with `graph LR` on line 1.

CONVERSATION:
""")


# =====================================================================
#  STEP 2 — Call OpenAI for intelligent extraction
# =====================================================================

import re


def check_mmdc_installed() -> bool:
    """Check if Mermaid CLI (mmdc) is installed and accessible."""
    return shutil.which("mmdc") is not None


def json_to_mermaid(json_str: str) -> str:
    """Convert JSON diagram structure to Mermaid syntax.
    
    Handles JSON diagrams returned by OpenAI when it ignores the 
    "return Mermaid code" instruction and returns JSON instead.
    """
    try:
        # Try to parse as JSON
        if isinstance(json_str, str):
            # Handle escaped JSON strings
            diagram_json = json.loads(json_str)
        else:
            diagram_json = json_str
            
        # Build Mermaid code from JSON structure
        lines = []
        graph_direction = diagram_json.get("graph", "LR")
        lines.append(f"graph {graph_direction}")
        
        # Mermaid reserved keywords that can't be used as style names
        reserved_keywords = {"start", "end", "graph", "subgraph", "class", "classDef"}
        
        # Helper function to sanitize labels for Mermaid
        def sanitize_label(text):
            """Make label safe for Mermaid syntax."""
            if not text:
                return "Step"
            # Remove or escape problematic characters
            text = str(text)
            text = text.replace('"', "'")  # Replace quotes with single quotes
            text = text.replace('&', 'and')
            text = text.replace('<', 'lt')
            text = text.replace('>', 'gt')
            return text[:50]  # Truncate to 50 chars for readability
        
        def make_safe_style_name(name):
            """Ensure style name is not a reserved keyword."""
            if name.lower() in reserved_keywords:
                return f"{name}Style"
            return name
        
        # Collect all nodes by ID for quick lookup
        nodes_by_id = {}
        for node in diagram_json.get("nodes", []):
            nodes_by_id[node.get("id")] = node
        
        # Create style name mapping
        style_name_map = {}
        for style_name in diagram_json.get("styles", {}).keys():
            safe_name = make_safe_style_name(style_name)
            style_name_map[style_name] = safe_name
        
        # Track which nodes are in subgraphs
        nodes_in_subgraphs = set()
        for sg in diagram_json.get("subgraphs", []):
            nodes_in_subgraphs.update(sg.get("nodes", []))
        
        # Add subgraphs first (swimlanes)
        subgraphs = diagram_json.get("subgraphs", [])
        if subgraphs:
            lines.append("")  # Blank line for readability
        
        for sg in subgraphs:
            sg_id = sg.get("id", "").replace(" ", "_")
            sg_label = sanitize_label(sg.get("label", sg_id))
            lines.append(f"  subgraph {sg_id}[\"{sg_label}\"]")
            
            # Add nodes in this subgraph
            sg_nodes = sg.get("nodes", [])
            for node_id in sg_nodes:
                node_data = nodes_by_id.get(node_id)
                if node_data:
                    label = sanitize_label(node_data.get("label", node_id))
                    node_type = node_data.get("type", "process")
                    node_class = node_data.get("class", "")
                    
                    # Format based on type
                    if node_type == "start":
                        lines.append(f"    {node_id}([{label}])")
                    elif node_type == "end":
                        lines.append(f"    {node_id}([{label}])")
                    elif node_type == "decision":
                        lines.append(f"    {node_id}{{{{{label}}}}}")
                    else:  # process
                        lines.append(f"    {node_id}[{label}]")
                    
                    # Add class if specified (using safe style name)
                    if node_class:
                        safe_class = style_name_map.get(node_class, node_class)
                        lines.append(f"    class {node_id} {safe_class}")
            
            lines.append("  end")
        
        # Add standalone nodes (not in subgraphs)
        standalone = [n for n in diagram_json.get("nodes", []) if n.get("id") not in nodes_in_subgraphs]
        if standalone:
            lines.append("")  # Blank line for readability
        
        for node in standalone:
            node_id = node.get("id", "")
            label = sanitize_label(node.get("label", node_id))
            node_type = node.get("type", "process")
            node_class = node.get("class", "")
            
            if node_type == "start":
                lines.append(f"  {node_id}([{label}])")
            elif node_type == "end":
                lines.append(f"  {node_id}([{label}])")
            elif node_type == "decision":
                lines.append(f"  {node_id}{{{{{label}}}}}")
            else:  # process
                lines.append(f"  {node_id}[{label}]")
            
            if node_class:
                safe_class = style_name_map.get(node_class, node_class)
                lines.append(f"  class {node_id} {safe_class}")
        
        # Add edges/connections
        edges = diagram_json.get("edges", [])
        if edges:
            lines.append("")  # Blank line for readability
        
        for edge in edges:
            from_id = edge.get("from", "")
            to_id = edge.get("to", "")
            label = edge.get("label", "")
            
            if not from_id or not to_id:
                continue
            
            if label:
                label = sanitize_label(label)
                lines.append(f"  {from_id} -->|{label}| {to_id}")
            else:
                lines.append(f"  {from_id} --> {to_id}")
        
        # Add style definitions at the end
        styles = diagram_json.get("styles", {})
        if styles:
            lines.append("")  # Blank line for readability
        
        for style_name, style_def in styles.items():
            safe_style_name = style_name_map.get(style_name, style_name)
            
            fill = style_def.get("fill", "#fff")
            stroke = style_def.get("stroke", "#000")
            stroke_width = style_def.get("stroke-width", "1px")
            color = style_def.get("color", "#000")
            
            lines.append(f"  classDef {safe_style_name} fill:{fill},stroke:{stroke},stroke-width:{stroke_width},color:{color}")
        
        return "\n".join(lines)
    
    except (json.JSONDecodeError, TypeError, AttributeError) as e:
        # If not valid JSON, return as-is (assume it's already Mermaid code)
        print(f"  Warning: Could not parse JSON diagram: {str(e)[:100]}")
        return json_str


def cleanup_mermaid(raw: str) -> str:
    """Ensure Mermaid code has proper line breaks and formatting.
    
    The AI sometimes returns everything on one line. This function
    forces proper multi-line formatting so mmdc can parse it.
    Also sanitizes special characters that break the Mermaid parser.
    """
    # Remove markdown fences
    text = raw.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
    if text.endswith("```"):
        text = text[:-3].strip()
    # Remove leading "mermaid" language tag
    if text.startswith("mermaid"):
        text = text[7:].strip()

    # Sanitize special characters inside node labels that break mmdc
    # Replace inside [...] and {...} labels only
    def sanitize_label(match):
        label = match.group(0)
        label = label.replace("€", " EUR")
        label = label.replace("£", " GBP")
        label = label.replace("$", " USD")
        label = label.replace("&", " and ")
        # < and > inside labels break Mermaid HTML parsing
        label = label.replace(" < ", " less than ")
        label = label.replace(" > ", " greater than ")
        label = label.replace("<", " less than ")
        label = label.replace(">", " greater than ")
        return label
    
    # Sanitize text inside square brackets [...] and curly braces {...}
    # but not the Mermaid arrow syntax -->
    text = re.sub(r'\[[^\]]+\]', sanitize_label, text)
    text = re.sub(r'\{[^}]+\}', sanitize_label, text)
    # Also sanitize pipe labels |...|
    text = re.sub(r'\|[^|]+\|', sanitize_label, text)

    # If it already has line breaks and looks well-formatted, return as-is
    if text.count("\n") > 5:
        return text

    # Otherwise, force line breaks at key Mermaid syntax points
    # These keywords/patterns should always start on a new line
    break_before = [
        r'subgraph ',
        r'end\b',
        r'classDef ',
        r'class ',
        r'style ',
        r'linkStyle ',
    ]

    # First, break before subgraph/end/classDef keywords
    for pattern in break_before:
        text = re.sub(r'(?<!\n)\s+(' + pattern + ')', r'\n\1', text)

    # Break after --> arrows and their labels: "A --> B" should end the line
    # Pattern: match node-->node or node-->|label|node, then ensure newline
    text = re.sub(r'(:::?\w+)(\s+)(?=[A-Z])', r'\1\n', text)

    # Break each arrow statement onto its own line
    # Match: "XN[...] --> " or "XN -->|...|" at positions that aren't at line start
    text = re.sub(r'(?<!\n)\s{2,}([A-Z]\w*[\[{(])', r'\n    \1', text)
    text = re.sub(r'(?<!\n)\s{2,}([A-Z]\w*\s*-->)', r'\n    \1', text)

    # If still mostly one line, do aggressive splitting
    if text.count("\n") < 5:
        # Split before any node definition that starts with a capital letter
        # followed by a digit and bracket
        parts = re.split(r'\s{2,}', text)
        lines = []
        indent = ""
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if part.startswith("subgraph"):
                lines.append(f"{indent}{part}")
                indent = "    "
            elif part == "end":
                indent = ""
                lines.append(f"{part}")
            elif part.startswith("graph "):
                lines.append(part)
            elif part.startswith("classDef") or part.startswith("class "):
                lines.append(part)
            else:
                lines.append(f"{indent}{part}")
        text = "\n".join(lines)

    return text.strip()


def build_conversation_text(session: dict) -> str:
    """Flatten the conversation history into readable text."""
    lines = []
    lines.append(f"Process Name: {session.get('project_name', 'Unknown')}")
    lines.append(f"Date: {session.get('start_time', 'Unknown')}")
    lines.append("=" * 60)

    for msg in session.get("conversation_history", []):
        role = msg.get("role", "unknown").upper()
        content = msg.get("content", "")
        # Strip any JSON status blocks from the content
        if "```json" in content:
            content = content[:content.index("```json")].strip()
        lines.append(f"\n{role}: {content}")

    return "\n".join(lines)


def call_openai(system_prompt: str, user_prompt: str, expect_json: bool = True) -> str | dict:
    """Call OpenAI and return the response (parsed as JSON if expected)."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("❌ ERROR: OPENAI_API_KEY not found in environment or .env file.")
        print("   Add it to your .env file: OPENAI_API_KEY=sk-...")
        sys.exit(1)

    client = OpenAI(api_key=api_key)

    print(f"   🤖 Calling {MODEL}...", end=" ", flush=True)
    response = client.chat.completions.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        temperature=0.2,  # Low temp for structured extraction
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    raw = response.choices[0].message.content.strip()
    print("Done ✓")

    if not expect_json:
        return raw

    # Clean and parse JSON
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]  # Remove first line
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"   ⚠️  JSON parse warning: {e}")
        print(f"   Raw response (first 500 chars): {raw[:500]}")
        # Try to fix common issues
        raw_fixed = raw.replace("'", '"')
        try:
            return json.loads(raw_fixed)
        except json.JSONDecodeError:
            return {"_raw": raw, "_error": str(e)}


def extract_all(session: dict) -> dict:
    """Run all extraction prompts and return combined results."""
    conversation_text = build_conversation_text(session)

    print("\n📊 STEP 1/4: Extracting SIPOC...")
    sipoc = call_openai(EXTRACTION_SYSTEM_PROMPT, SIPOC_PROMPT + conversation_text)

    print("📋 STEP 2/4: Extracting Process Map...")
    process_map = call_openai(EXTRACTION_SYSTEM_PROMPT, PROCESS_MAP_PROMPT + conversation_text)

    print("📈 STEP 3/4: Extracting Baseline Metrics...")
    baseline = call_openai(EXTRACTION_SYSTEM_PROMPT, BASELINE_PROMPT + conversation_text)

    print("🔀 STEP 4/4: Generating Mermaid Flowchart...")
    mermaid = call_openai(EXTRACTION_SYSTEM_PROMPT, MERMAID_PROMPT + conversation_text, expect_json=False)

    return {
        "sipoc": sipoc,
        "process_map": process_map,
        "baseline": baseline,
        "mermaid": mermaid,
        "project_name": session.get("project_name", "Process"),
        "date": session.get("start_time", datetime.now().isoformat()),
    }


# =====================================================================
#  STEP 3 — Generate Word Document
# =====================================================================

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def add_styled_table(doc, headers: list, rows: list, col_widths: list = None,
                     header_color: str = "1F4E79", header_text_color: str = "FFFFFF"):
    """Create a professional styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(header_text_color)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[1 + row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text) if cell_text else "")
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            # Alternate row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    return table


def add_heading_styled(doc, text: str, level: int = 1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.name = "Calibri"
    return heading


def generate_document(data: dict, output_path: str):
    """Generate the full Word document from extracted data."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    sipoc = data.get("sipoc", {})
    process_map = data.get("process_map", {})
    baseline = data.get("baseline", {})
    mermaid = data.get("mermaid", "")
    project_name = data.get("project_name", "Process")

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ════════════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Standardization Checkpoint")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = "Calibri"
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(sipoc.get("process_name", project_name))
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.name = "Calibri"

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"AS-IS Process Documentation\n"
                       f"Generated: {datetime.now().strftime('%B %d, %Y')}\n"
                       f"Process Owner: {sipoc.get('process_owner', 'TBD')}\n"
                       f"Frequency: {sipoc.get('frequency', 'TBD')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.font.name = "Calibri"

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)

    exec_summary = baseline.get("executive_summary", "")
    if not exec_summary or exec_summary == "Not discussed":
        exec_summary = (f"This document captures the AS-IS state of the "
                        f"{sipoc.get('process_name', project_name)} process "
                        f"as part of the Standardization Checkpoint in the "
                        f"Intelligent Automation Roadmap.")
    p = doc.add_paragraph(exec_summary)
    p.style.font.size = Pt(10)

    doc.add_paragraph("")

    # Key facts table
    key_facts = [
        ["Process Name", sipoc.get("process_name", project_name)],
        ["Process Owner", sipoc.get("process_owner", "TBD")],
        ["Frequency", sipoc.get("frequency", "TBD")],
        ["Organizational Position", sipoc.get("org_position", "TBD")],
        ["Systems & Tools", ", ".join(sipoc.get("systems_and_tools", ["TBD"]))],
    ]

    vol = baseline.get("volume", {})
    if isinstance(vol, dict) and vol.get("volume_per_period", "Not discussed") != "Not discussed":
        key_facts.append(["Volume", vol.get("volume_per_period", "TBD")])

    people = baseline.get("people", {})
    if isinstance(people, dict) and people.get("fte_involved", "Not discussed") != "Not discussed":
        key_facts.append(["People Involved", people.get("fte_involved", "TBD")])

    add_styled_table(doc, ["Attribute", "Details"], key_facts, col_widths=[5, 11])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  SIPOC DIAGRAM
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. SIPOC Diagram", level=1)

    p = doc.add_paragraph(
        "The SIPOC diagram provides a high-level overview of the process scope, "
        "identifying the key Suppliers, Inputs, Process steps, Outputs, and Customers."
    )

    # Build SIPOC rows — find the max length across all columns
    def format_items(items):
        if isinstance(items, list):
            result = []
            for item in items:
                if isinstance(item, dict):
                    name = item.get("name", "")
                    desc = item.get("description", "")
                    result.append(f"{name}: {desc}" if desc and desc != name else name)
                else:
                    result.append(str(item))
            return result
        return [str(items)]

    suppliers = format_items(sipoc.get("suppliers", []))
    inputs_ = format_items(sipoc.get("inputs", []))
    steps_hl = sipoc.get("process_steps_high_level", [])
    if isinstance(steps_hl, list):
        process_steps = [str(s) for s in steps_hl]
    else:
        process_steps = [str(steps_hl)]
    outputs = format_items(sipoc.get("outputs", []))
    customers = format_items(sipoc.get("customers", []))

    max_rows = max(len(suppliers), len(inputs_), len(process_steps), len(outputs), len(customers), 1)

    def pad_list(lst, length):
        return lst + [""] * (length - len(lst))

    suppliers = pad_list(suppliers, max_rows)
    inputs_ = pad_list(inputs_, max_rows)
    process_steps = pad_list(process_steps, max_rows)
    outputs = pad_list(outputs, max_rows)
    customers = pad_list(customers, max_rows)

    sipoc_rows = []
    for i in range(max_rows):
        sipoc_rows.append([suppliers[i], inputs_[i], process_steps[i], outputs[i], customers[i]])

    add_styled_table(
        doc,
        ["Suppliers", "Inputs", "Process (High-Level)", "Outputs", "Customers"],
        sipoc_rows,
        col_widths=[3.2, 3.2, 3.2, 3.2, 3.2],
    )

    # Known exceptions
    exceptions = sipoc.get("known_exceptions", [])
    if exceptions and exceptions != ["Not discussed"]:
        doc.add_paragraph("")
        add_heading_styled(doc, "Known Exceptions", level=2)
        for exc in exceptions:
            doc.add_paragraph(str(exc), style="List Bullet")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  DETAILED PROCESS MAP
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Detailed AS-IS Process Map", level=1)

    p = doc.add_paragraph(
        "The following table documents every step, decision point, and handoff "
        "in the current process as described during the analysis session."
    )

    steps = process_map.get("steps", [])
    if steps:
        step_rows = []
        for step in steps:
            pain = step.get("pain_point", "")
            if pain and pain != "null" and pain is not None:
                pain_text = f"⚠️ {pain}"
            else:
                pain_text = ""

            decision_text = ""
            if step.get("decision_point"):
                options = step.get("decision_options", [])
                decision_text = "🔀 Decision: " + " / ".join(options) if options else "🔀 Decision point"

            step_rows.append([
                step.get("step_number", ""),
                step.get("step_name", ""),
                step.get("description", ""),
                step.get("performer", ""),
                step.get("system_tool", ""),
                step.get("estimated_time", "") or "",
                (decision_text + "\n" + pain_text).strip(),
            ])

        add_styled_table(
            doc,
            ["#", "Step", "Description", "Performer", "System/Tool", "Time", "Notes"],
            step_rows,
            col_widths=[1, 2.5, 4, 2, 2, 1.5, 3],
        )
    else:
        doc.add_paragraph("No detailed process steps were extracted.", style="List Bullet")

    # Handoffs
    handoffs = process_map.get("handoffs", [])
    if handoffs:
        doc.add_paragraph("")
        add_heading_styled(doc, "3.1 Handoffs", level=2)
        handoff_rows = []
        for h in handoffs:
            handoff_rows.append([
                h.get("from_role", ""),
                h.get("to_role", ""),
                h.get("what", ""),
                h.get("how", ""),
            ])
        add_styled_table(doc, ["From", "To", "What", "How"], handoff_rows, col_widths=[3, 3, 5, 5])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  PAIN POINTS & IMPROVEMENT OPPORTUNITIES
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Pain Points & Improvement Opportunities", level=1)

    pain_points = process_map.get("pain_points_summary", [])
    if pain_points:
        pp_rows = []
        for pp in pain_points:
            affected = pp.get("affected_steps", [])
            affected_str = ", ".join(str(s) for s in affected) if affected else ""
            pp_rows.append([
                pp.get("category", ""),
                pp.get("description", ""),
                pp.get("impact", ""),
                affected_str,
            ])
        add_styled_table(
            doc,
            ["Category", "Pain Point", "Impact", "Affected Steps"],
            pp_rows,
            col_widths=[3, 5, 5, 3],
        )
    else:
        doc.add_paragraph("No specific pain points were identified during the session.")

    # Also list pain points from individual steps
    step_pains = [s for s in steps if s.get("pain_point") and s["pain_point"] != "null" and s["pain_point"] is not None]
    if step_pains:
        doc.add_paragraph("")
        add_heading_styled(doc, "4.1 Step-Level Pain Points", level=2)
        for sp in step_pains:
            doc.add_paragraph(
                f"Step {sp.get('step_number', '?')}: {sp.get('step_name', '')} — {sp.get('pain_point', '')}",
                style="List Bullet",
            )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  BASELINE METRICS
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Baseline Metrics", level=1)

    p = doc.add_paragraph(
        "The following metrics represent the current (AS-IS) performance of the process. "
        "These serve as the baseline against which improvements will be measured."
    )

    # Volume
    add_heading_styled(doc, "5.1 Volume & Frequency", level=2)
    vol = baseline.get("volume", {})
    if isinstance(vol, dict):
        vol_rows = [
            ["Frequency", vol.get("frequency", "Not discussed")],
            ["Volume per Period", vol.get("volume_per_period", "Not discussed")],
            ["Details", vol.get("details", vol.get("description", "Not discussed"))],
        ]
    else:
        vol_rows = [["Volume", str(vol)]]
    add_styled_table(doc, ["Metric", "Value"], vol_rows, col_widths=[5, 11])

    # Time
    doc.add_paragraph("")
    add_heading_styled(doc, "5.2 Time", level=2)
    time_data = baseline.get("time", {})
    if isinstance(time_data, dict):
        time_rows = [
            ["Total Cycle Time", time_data.get("total_cycle_time", "Not discussed")],
            ["Active Work Time", time_data.get("active_work_time", "Not discussed")],
            ["Waiting Time", time_data.get("waiting_time", "Not discussed")],
        ]
        add_styled_table(doc, ["Metric", "Value"], time_rows, col_widths=[5, 11])

        step_breakdown = time_data.get("step_breakdown", [])
        if step_breakdown:
            doc.add_paragraph("")
            sb_rows = []
            for sb in step_breakdown:
                sb_rows.append([
                    sb.get("step", ""),
                    sb.get("time", ""),
                    sb.get("notes", ""),
                ])
            add_styled_table(doc, ["Step", "Time", "Notes"], sb_rows, col_widths=[5, 3, 8])

    # People
    doc.add_paragraph("")
    add_heading_styled(doc, "5.3 People & Cost", level=2)
    ppl = baseline.get("people", {})
    if isinstance(ppl, dict):
        ppl_rows = [
            ["FTE Involved", ppl.get("fte_involved", "Not discussed")],
            ["Roles", ", ".join(ppl.get("roles", ["Not discussed"]))],
            ["Time Investment", ppl.get("time_investment_per_period", "Not discussed")],
            ["Cost Estimate", ppl.get("cost_estimate", "Not discussed")],
        ]
    else:
        ppl_rows = [["People", str(ppl)]]
    add_styled_table(doc, ["Metric", "Value"], ppl_rows, col_widths=[5, 11])

    # Quality
    doc.add_paragraph("")
    add_heading_styled(doc, "5.4 Quality", level=2)
    quality = baseline.get("quality", {})
    if isinstance(quality, dict):
        q_rows = [
            ["Error Rate", quality.get("error_rate", "Not discussed")],
            ["Rework Rate", quality.get("rework_rate", "Not discussed")],
            ["Common Errors", ", ".join(quality.get("common_errors", ["Not discussed"]))],
        ]
    else:
        q_rows = [["Quality", str(quality)]]
    add_styled_table(doc, ["Metric", "Value"], q_rows, col_widths=[5, 11])

    # Risk & Compliance
    doc.add_paragraph("")
    add_heading_styled(doc, "5.5 Risk & Compliance", level=2)
    risk = baseline.get("risk_and_compliance", {})
    if isinstance(risk, dict):
        r_rows = [
            ["Risks", ", ".join(risk.get("risks", ["Not discussed"]))],
            ["SLA / Deadlines", risk.get("sla", "Not discussed")],
            ["Compliance", risk.get("compliance_requirements", "Not discussed")],
        ]
    else:
        r_rows = [["Risk", str(risk)]]
    add_styled_table(doc, ["Metric", "Value"], r_rows, col_widths=[5, 11])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #  PROCESS FLOWCHART
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Process Flowchart", level=1)

    if mermaid and not isinstance(mermaid, dict):
        # Try to convert JSON diagram to Mermaid if needed
        mermaid_code = mermaid
        if mermaid.strip().startswith("{"):
            # Looks like JSON, try to convert it
            print(f"  ℹ️  Converting JSON diagram to Mermaid syntax...")
            mermaid_code = json_to_mermaid(mermaid)
        
        # Clean up and properly format the mermaid code
        mermaid_clean = cleanup_mermaid(mermaid_code)

        # Use ABSOLUTE paths — Windows + shell=True can fail with relative paths
        output_dir = os.path.abspath(os.path.dirname(output_path) or ".")
        base_name = os.path.splitext(os.path.basename(output_path))[0]

        mermaid_path = os.path.join(output_dir, f"{base_name}_flowchart.mmd")
        png_path = os.path.join(output_dir, f"{base_name}_flowchart.png")
        config_path = os.path.join(output_dir, f"{base_name}_mermaid_config.json")

        # Save .mmd file
        with open(mermaid_path, "w", encoding="utf-8") as f:
            f.write(mermaid_clean)
        print(f"  ✓ Mermaid source saved to: {mermaid_path}")

        # Check if mmdc is installed
        mmdc_available = check_mmdc_installed()
        if not mmdc_available:
            print("\n⚠️  WARNING: Mermaid CLI (mmdc) is not installed!")
            print("   The flowchart will be embedded as code text only, not as an image.")
            print("\n   To enable flowchart image rendering, install Mermaid CLI:")
            print("   ┌─────────────────────────────────────────────────────────┐")
            print("   │ npm install -g @mermaid-js/mermaid-cli                  │")
            print("   └─────────────────────────────────────────────────────────┘")
            print("   Then re-run this script.\n")

        # Try to render to PNG using Mermaid CLI (mmdc)
        flowchart_embedded = False

        if mmdc_available:
            try:
                print("  🎨 Rendering flowchart with Mermaid CLI...")

                # Create a config file for better rendering
                mmdc_config = {
                    "theme": "default",
                    "themeVariables": {
                        "fontSize": "14px",
                        "fontFamily": "Calibri, Arial, sans-serif"
                    }
                }
                with open(config_path, "w") as f:
                    json.dump(mmdc_config, f)

                # Build the mmdc command as a single string for shell execution
                mmdc_cmd = (
                    f'mmdc -i "{mermaid_path}" '
                    f'-o "{png_path}" '
                    f'-b white '
                    f'-w 2400 '
                    f'-H 1400 '
                    f'--configFile "{config_path}"'
                )

                print(f"     Running: {mmdc_cmd}")

                result = subprocess.run(
                    mmdc_cmd,
                    capture_output=True,
                    text=True,
                    timeout=60,
                    shell=True,  # Always use shell — ensures PATH is resolved
                )

                # Clean up config file
                if os.path.exists(config_path):
                    os.remove(config_path)

                print(f"     Return code: {result.returncode}")
                if result.stdout.strip():
                    print(f"     stdout: {result.stdout.strip()[:300]}")
                if result.stderr.strip():
                    print(f"     stderr: {result.stderr.strip()[:300]}")
                if os.path.exists(png_path):
                    file_size = os.path.getsize(png_path)
                    print(f"     PNG file size: {file_size} bytes")

                if os.path.exists(png_path) and os.path.getsize(png_path) > 0:
                    print(f"  ✓ Flowchart image rendered: {png_path} ({os.path.getsize(png_path)} bytes)")

                    # Add description
                    p = doc.add_paragraph(
                        "The diagram below shows the AS-IS process flow with all steps, "
                        "decision points, and handoffs as documented during the analysis session."
                    )

                    # For landscape flowcharts, switch to landscape page
                    # Add a landscape section just for the flowchart
                    new_section = doc.add_section()
                    new_section.orientation = WD_ORIENT.LANDSCAPE
                    new_section.page_width = Cm(29.7)
                    new_section.page_height = Cm(21.0)
                    new_section.left_margin = Cm(2.0)
                    new_section.right_margin = Cm(2.0)
                    new_section.top_margin = Cm(2.0)
                    new_section.bottom_margin = Cm(2.0)

                    # Available width in landscape: 29.7 - 4.0 = 25.7cm
                    landscape_width = 25.7

                    # Read image dimensions to maintain aspect ratio
                    try:
                        from PIL import Image as PILImage
                        img = PILImage.open(png_path)
                        img_width, img_height = img.size
                        aspect_ratio = img_height / img_width
                        img.close()

                        display_width_cm = landscape_width
                        display_height_cm = landscape_width * aspect_ratio

                        # If too tall for the page, scale down
                        max_height = 17.0  # cm, leaving room for margins
                        if display_height_cm > max_height:
                            display_height_cm = max_height
                            display_width_cm = max_height / aspect_ratio

                        doc.add_picture(png_path, width=Cm(display_width_cm))
                    except ImportError:
                        # Pillow not available, use fixed width
                        doc.add_picture(png_path, width=Cm(landscape_width))

                    # Center the image
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Switch back to portrait for remaining pages
                    back_to_portrait = doc.add_section()
                    back_to_portrait.orientation = WD_ORIENT.PORTRAIT
                    back_to_portrait.page_width = Cm(21.0)
                    back_to_portrait.page_height = Cm(29.7)
                    back_to_portrait.left_margin = Cm(2.5)
                    back_to_portrait.right_margin = Cm(2.5)
                    back_to_portrait.top_margin = Cm(2.0)
                    back_to_portrait.bottom_margin = Cm(2.0)

                    flowchart_embedded = True
                else:
                    print("  ✗ Flowchart rendering failed — PNG not created or empty")
                    print(f"     Expected path: {png_path}")

            except subprocess.TimeoutExpired:
                print("  ✗ Mermaid rendering timed out after 60 seconds.")
                print("   Try simplifying your flowchart or increasing the timeout.")
            except Exception as e:
                print(f"  ✗ Flowchart rendering error: {e}")
                import traceback
                traceback.print_exc()

        # If PNG rendering failed, include the Mermaid code as text fallback
        if not flowchart_embedded:
            p = doc.add_paragraph(
                "The flowchart could not be rendered as an image."
            )
            if not mmdc_available:
                doc.add_paragraph(
                    "The Mermaid CLI (mmdc) is not installed. "
                    "Install it with: npm install -g @mermaid-js/mermaid-cli"
                )
            else:
                doc.add_paragraph(
                    "The Mermaid.js code below can be rendered using mermaid.live, "
                    "VS Code Mermaid extension, or Confluence."
                )
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run(mermaid_clean)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    else:
        doc.add_paragraph("No flowchart was generated for this process.")

    # ════════════════════════════════════════════════════════════════
    #  APPENDIX — Raw Session Data
    # ════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "Appendix: Extracted Data (JSON)", level=1)

    p = doc.add_paragraph(
        "The following is the structured data extracted by the AI agent, "
        "preserved for reference and future processing."
    )

    # Save extracted data as JSON alongside the doc
    extracted_path = output_path.replace(".docx", "_extracted_data.json")
    with open(extracted_path, "w") as f:
        json.dump(data, f, indent=2, default=str)
    print(f"  ✓ Extracted data saved to: {extracted_path}")

    p = doc.add_paragraph(f"See: {Path(extracted_path).name}")
    run = p.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Save document ──
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    print(f"  ✓ Document saved to: {output_path}")


# =====================================================================
#  MAIN
# =====================================================================

def main():
    # Fix encoding on Windows to support emojis
    if sys.stdout.encoding.lower() != 'utf-8':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    
    print("=" * 60)
    print("  📄 INTELLIGENT DOCUMENT GENERATOR")
    print("  AI-Powered Process Documentation")
    print("=" * 60)

    # Determine session file
    if len(sys.argv) < 2:
        # Look for session files in outputs/
        outputs_dir = Path("outputs")
        if not outputs_dir.exists():
            print("\n❌ No outputs/ directory found.")
            print("   Usage: python intelligent_doc_generator.py <session_file.json>")
            sys.exit(1)

        session_files = sorted(outputs_dir.glob("*_session.json"), key=os.path.getmtime, reverse=True)
        if not session_files:
            print("\n❌ No session files found in outputs/")
            sys.exit(1)

        print(f"\n📁 Found {len(session_files)} session file(s):")
        for i, f in enumerate(session_files):
            print(f"   [{i + 1}] {f.name}")

        if len(session_files) == 1:
            session_path = session_files[0]
            print(f"\n   Using: {session_path.name}")
        else:
            choice = input(f"\n   Which file? [1-{len(session_files)}]: ").strip()
            try:
                session_path = session_files[int(choice) - 1]
            except (ValueError, IndexError):
                session_path = session_files[0]
                print(f"   Defaulting to: {session_path.name}")
    else:
        arg = sys.argv[1]
        session_path = Path(arg)
        if not session_path.exists():
            # Try adding outputs/ prefix and _session.json suffix
            session_path = Path("outputs") / f"{arg}_session.json"
        if not session_path.exists():
            print(f"\n❌ File not found: {arg}")
            print(f"   Also tried: {session_path}")
            sys.exit(1)

    # Load session
    print(f"\n📂 Loading: {session_path}")
    with open(session_path, "r") as f:
        session = json.load(f)

    project_name = session.get("project_name", "process")
    print(f"   Project: {project_name}")
    msg_count = len(session.get("conversation_history", []))
    print(f"   Messages in conversation: {msg_count}")

    # Extract data using AI
    print(f"\n🧠 Analyzing conversation with {MODEL}...")
    print("   (This makes 4 API calls — may take 30-60 seconds)\n")

    extracted_data = extract_all(session)

    # Generate document
    safe_name = project_name.replace(" ", "_").lower()
    output_path = f"outputs/{safe_name}_Standardization_Checkpoint.docx"

    print(f"\n📝 Generating Word document...")
    generate_document(extracted_data, output_path)

    print(f"\n{'=' * 60}")
    print(f"  ✅ DONE! Document generated successfully.")
    print(f"  📄 {output_path}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()