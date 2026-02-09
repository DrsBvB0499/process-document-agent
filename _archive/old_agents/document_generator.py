"""
Document Generator — Standardization Checkpoint
=================================================
Generates professional Word document (.docx) from the data
collected by the Process Analysis Agent.

Uses python-docx for document generation.

Dependencies:
    pip install python-docx Pillow
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


# ── Colors ──
class Colors:
    PRIMARY = RGBColor(0x1B, 0x4F, 0x72)
    SECONDARY = RGBColor(0x2E, 0x86, 0xC1)
    ACCENT = RGBColor(0xE7, 0x4C, 0x3C)
    SUCCESS = RGBColor(0x27, 0xAE, 0x60)
    TEXT_DARK = RGBColor(0x2C, 0x3E, 0x50)
    HEADER_BG = "1B4F72"
    ALT_ROW = "EBF5FB"
    SIPOC_BG = "D4E6F1"
    PAIN_BG = "FADBD8"
    GATE_BG = "D5F5E3"
    WHITE = "FFFFFF"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set cell margins (padding) in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_styled_cell(row, col_idx, text, bold=False, color=None, bg=None, align=None, size=10):
    """Add text to a table cell with styling."""
    cell = row.cells[col_idx]
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if bg:
        set_cell_shading(cell, bg)
    set_cell_margins(cell)
    return cell


class DocumentGenerator:
    """
    Generates Standardization Checkpoint Word documents from process data.
    """

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Aptos"
        font.size = Pt(11)
        font.color.rgb = Colors.TEXT_DARK

        # Heading styles
        for level, size, color in [
            ("Heading 1", 18, Colors.PRIMARY),
            ("Heading 2", 14, Colors.SECONDARY),
            ("Heading 3", 12, Colors.TEXT_DARK),
        ]:
            h_style = self.doc.styles[level]
            h_style.font.name = "Aptos"
            h_style.font.size = Pt(size)
            h_style.font.bold = True
            h_style.font.color.rgb = color

    def _add_table(self, headers, rows, col_widths=None, alt_row=True):
        """Add a formatted table with headers and data rows."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style header row
        for i, header in enumerate(headers):
            add_styled_cell(table.rows[0], i, header,
                            bold=True, color=RGBColor(255, 255, 255),
                            bg=Colors.HEADER_BG, size=10)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_data in enumerate(row_data):
                if isinstance(cell_data, dict):
                    add_styled_cell(table.rows[r_idx + 1], c_idx, **cell_data)
                else:
                    bg = Colors.ALT_ROW if (alt_row and r_idx % 2 == 1) else None
                    add_styled_cell(table.rows[r_idx + 1], c_idx, str(cell_data), bg=bg)

        # Set column widths if provided
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    row.cells[i].width = Cm(width)

        return table

    def generate(self, process_data: dict, flowchart_path: str, output_path: str) -> str:
        """
        Generate the complete Standardization Checkpoint document.

        Args:
            process_data: Dictionary with all process analysis data
            flowchart_path: Path to the process flow diagram PNG
            output_path: Where to save the .docx file

        Returns:
            The output file path
        """
        self._add_title_page(process_data)
        self._add_executive_summary(process_data)
        self._add_sipoc(process_data)
        self._add_process_map(process_data, flowchart_path)
        self._add_baseline(process_data)
        self._add_pain_points(process_data)
        self._add_process_statistics(process_data)
        self._add_gate_assessment(process_data)
        self._add_signoff()

        self.doc.save(output_path)
        return output_path

    def _add_title_page(self, data):
        """Add the title page."""
        # Spacing
        for _ in range(6):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("STANDARDIZATION CHECKPOINT")
        run.font.size = Pt(12)
        run.font.color.rgb = Colors.SECONDARY
        run.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get("process_name", "Process Name"))
        run.font.size = Pt(28)
        run.font.color.rgb = Colors.PRIMARY
        run.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get("process_description", ""))
        run.font.size = Pt(14)
        run.font.color.rgb = Colors.TEXT_DARK
        run.italic = True

        for _ in range(4):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Intelligent Automation Roadmap — STOP Gate 1")
        run.font.size = Pt(11)
        run.font.color.rgb = Colors.ACCENT
        run.bold = True

        info = [
            f"Process Owner: {data.get('process_owner', 'TBD')}",
            f"Department: {data.get('department', 'TBD')}",
            f"Date: {datetime.now().strftime('%B %Y')}",
        ]
        for line in info:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("✓ READY FOR OPTIMIZATION")
        run.font.size = Pt(12)
        run.font.color.rgb = Colors.SUCCESS
        run.bold = True

        self.doc.add_page_break()

    def _add_executive_summary(self, data):
        """Add executive summary section."""
        self.doc.add_heading("Executive Summary", level=1)
        self.doc.add_paragraph(data.get("executive_summary", ""))
        self.doc.add_page_break()

    def _add_sipoc(self, data):
        """Add SIPOC analysis section."""
        self.doc.add_heading("SIPOC Analysis", level=1)
        self.doc.add_paragraph(
            "The SIPOC framework provides a high-level view of the process scope, "
            "identifying all key stakeholders and the flow of value from suppliers to customers."
        )

        sipoc = data.get("sipoc", {})
        sipoc_rows = [
            [{"text": "Suppliers", "bold": True, "bg": Colors.SIPOC_BG, "size": 10},
             sipoc.get("suppliers", "")],
            [{"text": "Inputs", "bold": True, "bg": Colors.SIPOC_BG, "size": 10},
             sipoc.get("inputs", "")],
            [{"text": "Process", "bold": True, "bg": Colors.SIPOC_BG, "size": 10},
             sipoc.get("process", "")],
            [{"text": "Outputs", "bold": True, "bg": Colors.SIPOC_BG, "size": 10},
             sipoc.get("outputs", "")],
            [{"text": "Customers", "bold": True, "bg": Colors.SIPOC_BG, "size": 10},
             sipoc.get("customers", "")],
        ]
        self._add_table(["SIPOC Analysis", ""], sipoc_rows, col_widths=[4, 14])

        self.doc.add_paragraph()
        self.doc.add_heading("Process Context", level=2)

        context = data.get("context", {})
        context_rows = [
            [{"text": k, "bold": True, "bg": Colors.SIPOC_BG, "size": 10}, v]
            for k, v in context.items()
        ]
        self._add_table(["Field", "Value"], context_rows, col_widths=[5, 13])
        self.doc.add_page_break()

    def _add_process_map(self, data, flowchart_path):
        """Add AS-IS process map section with embedded diagram."""
        self.doc.add_heading("AS-IS Process Map", level=1)
        self.doc.add_paragraph(
            "The following table documents each step in the current process, "
            "including the performer, systems used, decision points, "
            "identified pain points, and estimated time per step."
        )

        steps = data.get("process_steps", [])
        headers = ["#", "Step", "Performer", "System", "Decision Point", "Pain Point", "Time"]
        self._add_table(headers, steps, col_widths=[1, 4.5, 2.5, 2.5, 2.5, 3, 2])

        # Embed flowchart
        self.doc.add_paragraph()
        self.doc.add_heading("Process Flow Diagram", level=2)
        self.doc.add_paragraph(
            "The following diagram provides a visual representation of the complete "
            "AS-IS process flow, showing all steps, decision points, exception paths, "
            "and the handoff between systems and team members."
        )

        if os.path.exists(flowchart_path):
            self.doc.add_picture(flowchart_path, width=Inches(6.2))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.doc.add_paragraph(
                f"[Flowchart image not found at {flowchart_path}]"
            )

        self.doc.add_page_break()

    def _add_baseline(self, data):
        """Add baseline measurement section."""
        self.doc.add_heading("Baseline Measurement", level=1)
        self.doc.add_paragraph(
            "The baseline establishes the current cost of the process in time, effort, "
            "and risk. These metrics serve as the benchmark against which optimization "
            "and automation improvements will be measured."
        )

        metrics = data.get("baseline_metrics", [])
        self._add_table(["Metric", "Value", "Notes"], metrics, col_widths=[6, 5, 7])

        self.doc.add_paragraph()
        summary = data.get("baseline_summary", "")
        if summary:
            p = self.doc.add_paragraph()
            run = p.add_run(summary)
            run.bold = True

        self.doc.add_page_break()

    def _add_pain_points(self, data):
        """Add pain points section."""
        self.doc.add_heading("Identified Pain Points", level=1)
        self.doc.add_paragraph(
            "The following pain points were identified during the process analysis "
            "and are prioritized by severity."
        )

        pains = data.get("pain_points", [])
        self._add_table(
            ["#", "Pain Point", "Description", "Impact", "Severity"],
            pains, col_widths=[1, 4, 4.5, 4, 2.5]
        )
        self.doc.add_page_break()

    def _add_process_statistics(self, data):
        """Add process statistics summary."""
        self.doc.add_heading("Process Statistics Overview", level=1)

        stats = data.get("statistics", [])
        self._add_table(["Metric", "Value"], stats, col_widths=[8, 10])
        self.doc.add_page_break()

    def _add_gate_assessment(self, data):
        """Add standardization gate assessment."""
        self.doc.add_heading("Standardization Gate Assessment", level=1)
        self.doc.add_paragraph(
            "The following table evaluates whether the process meets all criteria "
            "required to pass through STOP Gate 1 (Standardization) of the "
            "Intelligent Automation Roadmap."
        )

        criteria = data.get("gate_criteria", [])
        self._add_table(
            ["Checkpoint Criteria", "Status", "Evidence"],
            criteria, col_widths=[6, 3, 9]
        )

        # Gate verdict
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("✓  STANDARDIZATION GATE: PASSED")
        run.font.size = Pt(16)
        run.font.color.rgb = Colors.SUCCESS
        run.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("This process is approved to proceed to the Optimization phase.")
        run.italic = True

        # Next steps
        self.doc.add_paragraph()
        self.doc.add_heading("Recommended Next Steps", level=2)
        for step in data.get("next_steps", []):
            self.doc.add_paragraph(step, style="List Bullet")

        self.doc.add_page_break()

    def _add_signoff(self):
        """Add sign-off section."""
        self.doc.add_heading("Sign-off", level=1)
        self.doc.add_paragraph(
            "By signing below, stakeholders confirm that the Standardization phase "
            "is complete and the process may proceed to Optimization."
        )
        self.doc.add_paragraph()

        roles = ["Process Owner", "Business Analyst", "Team Lead"]
        self._add_table(
            ["Role", "Name", "Signature / Date"],
            [[r, "", ""] for r in roles],
            col_widths=[5, 6, 7]
        )


if __name__ == "__main__":
    print("DocumentGenerator class ready for use.")
    print("Use session_to_document.py to generate documents from session files.")
